from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as Docx

# 获取带有元数据的pdf文件内容
def get_pdf_text_with_metadata(pdf_path):
    pdf_reader = PdfReader(pdf_path)
    documents = [
        Document(
            page_content=page.extract_text(),
            metadata={
                "source": pdf_path,
                "page": i+1
            }
        )
        for i, page in enumerate(pdf_reader.pages)
    ]
    return documents


def get_docx_text_with_metadata(docx_path):
    doc = Docx(docx_path)
    documents = [
        Document(
            page_content=para.text,
            metadata={
                "source": docx_path,
                "paragraphs": i+1
            }
        )
        for i, para in enumerate(doc.paragraphs)
    ]
    return documents


# 拆分带有元数据的文档
def split_documents_with_metadata(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=300,
        separators=["\n## ", "\n\n", "\n", "。", "！"]
    )
    chunks = text_splitter.split_documents(documents)  # 关键：使用 split_documents 而非 split_text
    return chunks


# 获取pdf文件内容
def get_pdf_text(pdf_path):
    text = ""
    pdf_reader = PdfReader(pdf_path)
    for page in pdf_reader.pages:
        text += page.extract_text()

    return text

# 拆分文本
def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        # chunk_size=768,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks