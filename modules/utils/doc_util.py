import os

from PyPDF2 import PdfReader
from docx import Document as Docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


class DocUtil:
    # 获取带有元数据的pdf文件内容
    @staticmethod
    def __get_pdf_text(pdf_path):
        pdf_reader = PdfReader(pdf_path)
        documents = [
            Document(
                page_content=page.extract_text(),
                metadata={
                    "source": pdf_path,
                    "page": i+1
                }
            )
            for i, page in enumerate(pdf_reader.pages)
        ]
        return documents

    @staticmethod
    def __get_docx_text(docx_path):
        doc = Docx(docx_path)
        documents = [
            Document(
                page_content=para.text,
                metadata={
                    "source": docx_path,
                    "paragraphs": i+1
                }
            )
            for i, para in enumerate(doc.paragraphs)
        ]
        return documents


    # 拆分带有元数据的文档
    @staticmethod
    def __split_documents(documents):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=300,
            separators=["\n## ", "\n\n", "\n", "。", "！"]
        )
        chunks = text_splitter.split_documents(documents)  # 关键：使用 split_documents 而非 split_text
        return chunks

    # 统一处理文件的方法
    @staticmethod
    def process_file(file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if file_path.endswith(".pdf"):
            documents = DocUtil.__get_pdf_text(file_path)
        elif file_path.endswith(".docx"):
            documents = DocUtil.__get_docx_text(file_path)
        else:
            raise ValueError("不支持的文件类型")

        # 对文档进行分块
        chunks = DocUtil.__split_documents(documents)
        return chunks